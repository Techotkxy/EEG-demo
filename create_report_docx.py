"""Generate EEG Viewer Report as .docx"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('OpenBCI EEG Viewer – Technical Report', 0)
doc.add_paragraph('Document: Cogwear Demo EEG Viewer')
doc.add_paragraph('Version: Multi-headband real-time viewer with synchrony analysis')

doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'The EEG Viewer is a real-time visualization and analysis application for OpenBCI Cyton headbands. '
    'It supports 1–4 headbands simultaneously (4–16 channels), with preprocessing, frequency-band decomposition, '
    'and inter-subject EEG synchrony (Phase Locking Value) when multiple headbands are connected.'
)

doc.add_heading('2. Hardware Support', level=1)
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Feature'
table.rows[0].cells[1].text = 'Specification'
table.rows[1].cells[0].text = 'Headbands'
table.rows[1].cells[1].text = '1–4 OpenBCI Cyton headbands'
table.rows[2].cells[0].text = 'Channels per headband'
table.rows[2].cells[1].text = '4 (AF7, FP1, FP2, AF8)'
table.rows[3].cells[0].text = 'Total channels'
table.rows[3].cells[1].text = '4, 8, 12, or 16'
table.rows[4].cells[0].text = 'Sample rate'
table.rows[4].cells[1].text = '250 Hz'
table.rows[5].cells[0].text = 'Connection'
table.rows[5].cells[1].text = 'USB dongle per headband → COM port (e.g. COM3, COM4)'
table.rows[6].cells[0].text = 'Radio'
table.rows[6].cells[1].text = 'Each headband must use a different radio channel (see RADIO_SETUP.md)'

doc.add_heading('3. Usage', level=1)
table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Table Grid'
table2.rows[0].cells[0].text = 'Command'
table2.rows[0].cells[1].text = 'Description'
table2.rows[1].cells[0].text = 'python view_raw_signal.py COM3'
table2.rows[1].cells[1].text = 'Single headband'
table2.rows[2].cells[0].text = 'python view_raw_signal.py COM3 COM4'
table2.rows[2].cells[1].text = 'Two headbands'
table2.rows[3].cells[0].text = 'python view_raw_signal.py COM3 COM4 COM5 COM6'
table2.rows[3].cells[1].text = 'Four headbands (default when no ports specified)'

doc.add_heading('4. Display Windows', level=1)
doc.add_paragraph('Raw EEG – Preprocessed signal (0.3–50 Hz bandpass, notch, ICA, z-score) in 5-second rolling window', style='List Bullet')
doc.add_paragraph('Theta (4–8 Hz) – Band-filtered theta activity', style='List Bullet')
doc.add_paragraph('Alpha (8–13 Hz) – Band-filtered alpha activity', style='List Bullet')
doc.add_paragraph('Beta (13–30 Hz) – Band-filtered beta activity', style='List Bullet')
doc.add_paragraph('Gamma (30–60 Hz) – Band-filtered gamma activity', style='List Bullet')
doc.add_paragraph('EEG Synchrony (PLV) – Pairwise Phase Locking Value in alpha band (8–13 Hz); shown only when 2+ headbands are connected', style='List Bullet')
doc.add_paragraph('Plots are arranged in a grid (2×2 for 4 channels, 2×4 for 8, 4×4 for 16). Packet loss and flat-signal indicators appear at the bottom of the Raw EEG window.')

doc.add_heading('5. Signal Processing Pipeline', level=1)
doc.add_heading('5.1 Preprocessing (applied to all channels)', level=2)
table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Table Grid'
table3.rows[0].cells[0].text = 'Step'
table3.rows[0].cells[1].text = 'Setting'
table3.rows[0].cells[2].text = 'Value'
table3.rows[1].cells[0].text = 'Bandpass'
table3.rows[1].cells[1].text = 'Low / High cutoff'
table3.rows[1].cells[2].text = '0.3 Hz / 50 Hz (4th-order Butterworth)'
table3.rows[2].cells[0].text = 'Notch'
table3.rows[2].cells[1].text = 'Frequency / Q'
table3.rows[2].cells[2].text = '50 Hz (or 60 Hz) / Q=30'
table3.rows[3].cells[0].text = 'Common Average Reference (CAR)'
table3.rows[3].cells[1].text = 'Per headband'
table3.rows[3].cells[2].text = 'Subtract mean of 4 channels from each channel within that headband'
table3.rows[4].cells[0].text = 'ICA'
table3.rows[4].cells[1].text = 'FastICA, kurtosis-based rejection'
table3.rows[4].cells[2].text = 'Reject components with |z| > 1.96; refit every 2.5 s'

doc.add_heading('5.2 Display-Only Processing', level=2)
doc.add_paragraph('These steps affect visualization only, not stored or exported data:', style='List Bullet')
doc.add_paragraph('Median filter (5-point) to reduce spikes', style='List Bullet')
doc.add_paragraph('Outlier rejection: replace values beyond ±2 z-score with local median', style='List Bullet')
doc.add_paragraph('Moving average smoothing (15-point)', style='List Bullet')
doc.add_paragraph('Display clipping at ±2.5 z-score', style='List Bullet')
doc.add_paragraph('Y-axis auto-range', style='List Bullet')

doc.add_heading('6. EEG Synchrony (Multi-Headband)', level=1)
doc.add_paragraph(
    'When 2 or more headbands are connected: Phase Locking Value (PLV) is computed in the alpha band (8–13 Hz) '
    'for all pairwise combinations. Each headband\'s signal is the mean of its 4 channels. '
    'PLV ranges from 0 (no phase locking) to 1 (strong phase locking). Bar chart updates in real time (~30 Hz).'
)

doc.add_heading('7. Diagnostics and Monitoring', level=1)
doc.add_paragraph('Packet loss – Per-headband count and percentage', style='List Bullet')
doc.add_paragraph('Flat-signal indicator – "⚠flat" when a headband\'s signal is much flatter than others', style='List Bullet')
doc.add_paragraph('Cross-talk check – Warns if two headbands show correlation > 0.92 (possible same radio channel)', style='List Bullet')
doc.add_paragraph('Radio channel reminder – Startup message when 2+ headbands; see RADIO_SETUP.md', style='List Bullet')

doc.add_heading('8. Configuration Summary', level=1)
table4 = doc.add_table(rows=15, cols=2)
table4.style = 'Table Grid'
configs = [
    ('LOW_CUTOFF', '0.3 Hz'),
    ('HIGH_CUTOFF', '50 Hz'),
    ('NOTCH_HZ', '50 (or 60 for Americas)'),
    ('NOTCH_Q', '30'),
    ('USE_CAR', 'True'),
    ('ICA_DISABLED', 'False'),
    ('Z_THRESHOLD', '1.96'),
    ('ICA_UPDATE_INTERVAL_MS', '2500'),
    ('PLOT_REFRESH_MS', '33 (~30 Hz)'),
    ('DISPLAY_SMOOTH', '15'),
    ('OUTLIER_REJECT_Z', '2.0'),
    ('DISPLAY_CLIP_Z', '2.5'),
    ('DISPLAY_MEDIAN_FILTER', '5'),
    ('SYNC_BAND (PLV)', '8–13 Hz (alpha)'),
]
for i, (k, v) in enumerate(configs):
    table4.rows[i].cells[0].text = k
    table4.rows[i].cells[1].text = v

doc.add_heading('9. Dependencies', level=1)
doc.add_paragraph('Python 3, pyqtgraph, PyQt5, numpy, scipy, scikit-learn, pyserial', style='List Bullet')

doc.add_heading('10. File Structure', level=1)
doc.add_paragraph('view_raw_signal.py – Main viewer application', style='List Bullet')
doc.add_paragraph('lib/open_bci_v3.py – OpenBCI serial protocol and board interface', style='List Bullet')
doc.add_paragraph('RADIO_SETUP.md – Radio channel configuration guide', style='List Bullet')

doc.save('EEG_Viewer_Report.docx')
print('Created EEG_Viewer_Report.docx')
